from openai import OpenAI
import PyPDF2 as pdf
import json
from docx import Document
import re


def configure_openai(api_key: str):
    """
    Create and return an OpenAI client.
    """
    try:
        client = OpenAI(api_key=api_key)
        return client
    except Exception as e:
        raise Exception(f"Failed to configure OpenAI API: {str(e)}")


def get_openai_response(client, prompt: str):
    """
    Generate ATS evaluation using OpenAI with strict JSON validation.
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini", 
            messages=[
                {
                    "role": "system",
                    "content": "You are a strict ATS resume evaluation engine. Respond ONLY with valid JSON."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.2
        )

        if not response or not response.choices:
            raise Exception("Received empty response from OpenAI.")

        raw_text = response.choices[0].message.content.strip()

        try:
            response_json = json.loads(raw_text)
        except json.JSONDecodeError:
            # ---- Extract JSON if model adds text ---- #
            match = re.search(r"\{.*\}", raw_text, re.DOTALL)
            if not match:
                raise Exception("Could not extract valid JSON from OpenAI response.")
            response_json = json.loads(match.group())

        # ---- Validate required fields ---- #
        required_fields = ["JD Match", "Missing Keywords", "Profile Summary"]
        for field in required_fields:
            if field not in response_json:
                raise ValueError(f"Missing required field: {field}")

        return response_json

    except Exception as e:
        raise Exception(f"Error generating response: {str(e)}")
    


def extract_pdf_text(uploaded_file):
    """
    Extract text from uploaded PDF with error handling.
    """
    try:
        reader = pdf.PdfReader(uploaded_file)

        if not reader.pages:
            raise Exception("PDF file is empty.")

        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)

        if not text:
            raise Exception("No text could be extracted from the PDF.")

        return " ".join(text)

    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


def extract_docx_text(uploaded_file):
    """
    Robust DOCX text extraction for resumes.
    Handles paragraphs, tables, and formatted layouts.
    """
    try:
        doc = Document(uploaded_file)
        text = []

        for para in doc.paragraphs:
            content = para.text.strip()
            if content:
                text.append(content)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text.append(cell_text)

        combined_text = " ".join(text).strip()

        if not combined_text:
            raise Exception(
                "The document contains text in unsupported formats (images or shapes)."
            )

        return combined_text

    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")

def prepare_prompt(resume_text: str, job_description: str):
    """
    Prepare ATS evaluation prompt.
    """
    if not resume_text or not job_description:
        raise ValueError("Resume text and job description must not be empty.")

    prompt_template = f"""
You are an expert ATS (Applicant Tracking System) specialist with deep expertise in:
- Software Engineering
- Data Science
- Data Analysis
- Big Data Engineering
- Machine Learning

Evaluate the resume against the job description considering a highly competitive job market.

Resume:
{resume_text}

Job Description:
{job_description}

Respond ONLY in the following JSON format:
{{
  "JD Match": "percentage between 0-100",
  "Missing Keywords": ["keyword1", "keyword2"],
  "Profile Summary": "Detailed ATS-based evaluation and improvement suggestions"
}}
"""
    return prompt_template.strip()
